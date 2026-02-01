#!/usr/bin/env python3
"""Generate resume docx matching the 이력서 양식(엔지니어) template."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')


def add_section_heading(text):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # Add bottom border
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): '999999',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_bullet(text, bold_prefix=None, indent_level=0):
    p = doc.add_paragraph()
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    fmt = p.paragraph_format
    fmt.left_indent = Cm(0.5 + indent_level * 0.5)
    fmt.first_line_indent = Cm(-0.3)

    run = p.add_run('• ')
    run.font.size = Pt(10)

    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(10)

    run_t = p.add_run(text)
    run_t.font.size = Pt(10)
    return p


def add_normal(text, bold=False, size=10, space_after=2):
    p = doc.add_paragraph()
    p.space_before = Pt(1)
    p.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def add_bold_and_normal(bold_text, normal_text, size=10):
    p = doc.add_paragraph()
    p.space_before = Pt(1)
    p.space_after = Pt(2)
    run_b = p.add_run(bold_text)
    run_b.bold = True
    run_b.font.size = Pt(size)
    run_n = p.add_run(normal_text)
    run_n.font.size = Pt(size)
    return p


# ============================================================
# 1. 개인정보
# ============================================================
add_section_heading('개인정보')
add_bold_and_normal('이    름 : ', '정주안 (Juahn Jeong) / 남')
add_bold_and_normal('생년월일 : ', '(직접 기입)')
add_bold_and_normal('현재연봉 : ', '(직접 기입)')
add_bold_and_normal('희망연봉 : ', '(직접 기입)')
add_bold_and_normal('입사시기 : ', '(직접 기입)')
add_bold_and_normal('연 락 처 : ', '(직접 기입)')
add_bold_and_normal('이 메 일 : ', 'johnnyapu15@gmail.com')


# ============================================================
# 2. 학력사항
# ============================================================
add_section_heading('학력사항')
add_normal('2018.03~2019.08\t서울시립대학교 대학원 컴퓨터과학과 석사 졸업 (인공지능 시스템 성능 개선 연구 / 학·석사 연계 조기 취득)')
add_normal('2012.03~2017.02\t서울시립대학교 컴퓨터과학부 학사 졸업')


# ============================================================
# 3. 경력사항
# ============================================================
add_section_heading('경력사항 (총 경력 : 5년 7개월)')
add_normal('2022.08~재직 중\tBucketplace (오늘의집) / XR Engineering / Backend Engineer / Software Engineer')
add_normal('2020.01~2022.07\tLG CNS / DX사업부 / Backend Developer / Consultant')


# ============================================================
# 4. 핵심역량
# ============================================================
add_section_heading('핵심역량')
add_bullet('Kotlin/Spring Boot 기반 백엔드 아키텍처 설계부터 운영까지 E2E로 담당. 트래픽 10배 증가 환경에서 트랜잭션 안정성과 서비스 가용성(99.95%) 확보')
add_bullet('장애 전파 방지 구조(Circuit Breaker), 비동기 처리, Kafka 기반 이벤트 드리븐 아키텍처로 분산 시스템의 신뢰성 확보')
add_bullet('Datadog 기반 모니터링 체계 구축 및 Custom Metric 기반 오토스케일링으로 운영 병목 사전 대응')
add_bullet('자동화 파이프라인으로 운영 효율 27% 향상, 연간 40일 절감 등 생산성에 지속 기여하며 매 분기 Top Contributor 선정')
add_bullet('CI/CD 파이프라인 구축·개선, 테스트 자동화 주도로 배포 안정성 및 품질 확보')


# ============================================================
# 5. 기술스택
# ============================================================
add_section_heading('기술스택')
add_bullet('Kotlin, Java, Python, TypeScript', bold_prefix='Languages: ')
add_bullet('Spring Boot, Node.js, React', bold_prefix='Frameworks: ')
add_bullet('AWS, Kubernetes, Kafka, MySQL, MongoDB, Redis, Elasticsearch', bold_prefix='Backend/Infra: ')
add_bullet('Git, Docker, ArgoCD, Datadog, Jenkins', bold_prefix='Tools: ')


# ============================================================
# 6. 경력 세부사항
# ============================================================
add_section_heading('경력 세부사항')

# --- Bucketplace ---
p = doc.add_paragraph()
p.space_before = Pt(6)
p.space_after = Pt(4)
run = p.add_run('2022.08~ 재직 중    Bucketplace (오늘의집) / XR Engineering / Backend Engineer / Software Engineer')
run.bold = True
run.font.size = Pt(10)

add_bold_and_normal('[회사소개]', '')
add_normal('- 업종 및 제품 : 인테리어 플랫폼 / 커머스·콘텐츠·3D/AR 서비스 운영')
add_normal('- 매출액: 3,891억원 (2023년 기준) / 직원 수: 약 1,000명')

add_bold_and_normal('\n[담당 업무]', '')
add_normal('1. Kotlin/Spring Boot 기반 글로벌 서비스의 백엔드 아키텍처 설계, 구축 및 운영')
add_normal('2. 서비스 안정성·가용성 확보 및 모니터링 체계 구축')
add_normal('3. 자동화 파이프라인 설계 및 구축을 통한 운영 효율화')

add_bold_and_normal('\n[주요 업무/성과]', '')

# AI Interior Service
add_normal('1. AI 인테리어 서비스 런칭 및 고도화 (Kotlin/Spring Boot)', bold=True)
add_bullet('6주 안에 글로벌 런칭이 필요했고, 외부 API의 불안정성으로 서비스 장애 위험이 있는 상황')
add_bullet('멀티 모듈 구조로 API/워커를 분리 배포하고, Kafka 파티션 키로 사용자별 요청 순서를 보장하며 Kafka lag 기반 오토스케일링으로 안정적인 처리량 유지')
add_bullet('Retry, Circuit Breaker로 외부 API 장애를 격리하고, Provider 추상화로 장애 시 빠른 전환이 가능한 구조를 설계하여 4주 만에 런칭')
add_bullet('Datadog Custom Metric 기반 모니터링 체계를 구축하여 장애·성능 이슈를 사전에 탐지하고 대응')
add_bullet('Apple/Google 결제 라이프사이클을 비동기 웹훅과 상태 머신으로 통합하여 IAP 구독 모델을 일주일 만에 구축')
add_bullet('4주 만에 170개국 글로벌 런칭 완료. 외부 장애에도 안정적 처리, 구독 모델 검증 및 시스템 안정성 확보', bold_prefix='→ ')

# 3D Pipeline
add_normal('\n2. 3D 파이프라인 자동화', bold=True)
add_bullet('수작업 기반 파이프라인의 관리·최적화·생성 프로세스가 한계에 부딪혀 있는 상황')
add_bullet('상태 머신으로 저장→최적화→배포를 자동화하고, stateless 워커와 DB 기반 진행 관리로 장애 복구와 수평 확장이 가능하게 구축')
add_bullet('반복 실험으로 최적 파라미터를 찾고, fallback 로직으로 예외 케이스에 대응해 품질 유지')
add_bullet('운영 전용 대시보드를 구축하여 엔지니어 의존 없이 운영팀이 직접 파이프라인 관리 가능하게 함')
add_bullet('GPT-4o 기반 이미지 품질 점수화 → 3D 모델 생성까지 E2E 자동화 파이프라인 구축')
add_bullet('GPU 20%↓, 파일 크기 51%↓, 연간 40일 이상 절감 → Top Contributor 선정', bold_prefix='→ ')
add_bullet('생산량 1,176%↑, 비용 88%↓ → Eng Award 수상', bold_prefix='→ ')

# 3D Interior Design Service
add_normal('\n3. 서비스 운영 및 신뢰성 확보 (Room Planner)', bold=True)
add_bullet('트래픽 10.2배 증가에도 99.95% 가용성을 확보하며 WAU 704% 성장, 연간 GMV 600%(8.6억) 달성에 기여')
add_bullet('장애·성능 분석 및 재발 방지를 위한 구조 개선 주도')

# AI Native Culture
add_normal('\n4. 배포·테스트 자동화 및 AI 활용 문화 선도', bold=True)
add_bullet('릴리즈 노트 자동 생성, Unit Test 고도화, JMeter 기반 부하 테스트 설계 가이드 수립으로 배포 안정성·품질 확보')
add_bullet('AI Native MVP 및 AI Award 수상', bold_prefix='→ ')

add_bold_and_normal('\n이직사유 : ', '(직접 기입)')

# --- LG CNS ---
p = doc.add_paragraph()
p.space_before = Pt(16)
p.space_after = Pt(4)
run = p.add_run('2020.01~2022.07    LG CNS / DX사업부 / Backend Developer / Consultant')
run.bold = True
run.font.size = Pt(10)

add_bold_and_normal('[회사소개]', '')
add_normal('- 업종 및 제품 : IT 서비스 / SI·SM·클라우드·AI 사업')
add_normal('- 매출액: 약 5조원 (2022년 기준) / 직원 수: 약 7,000명')

add_bold_and_normal('\n[담당 업무]', '')
add_normal('1. 실시간 싱크 서버 아키텍처 설계 및 운영')
add_normal('2. CI/CD 파이프라인 구축 및 배포·테스트 자동화')

add_bold_and_normal('\n[주요 업무/성과]', '')
add_bullet('Node.js · Redis 기반 실시간 싱크 서버 아키텍처를 설계하고, 다수 동시 접속 환경에서 동시성 이슈 없이 안정적으로 운영')
add_bullet('GitLab/Jenkins CI/CD 파이프라인 구축으로 빌드·테스트 자동화 및 배포 리드타임 단축')

add_bold_and_normal('\n이직사유 : ', '(직접 기입)')


# ============================================================
# 7. 기타사항
# ============================================================
add_section_heading('기타사항')
add_bold_and_normal('병  역 : ', '(직접 기입)')
add_bold_and_normal('어  학 : ', '(직접 기입)')
add_bold_and_normal('자격증 : ', '(직접 기입)')


# ============================================================
# 8. 자기소개
# ============================================================
add_section_heading('자기소개')

intro = """Kotlin/Spring Boot 기반 백엔드 아키텍처를 설계하고, 서비스의 런칭부터 안정화, 확장까지 E2E로 담당해온 백엔드 엔지니어입니다.

서비스 안정성을 최우선으로 설계해왔습니다. 외부 API가 불안정한 환경에서도 Circuit Breaker, 비동기 처리, Kafka 기반 이벤트 드리븐 아키텍처로 장애 전파를 방지하고 트래픽 10배 증가 속에서도 99.95%의 가용성을 유지했습니다. 또한 Datadog 기반 모니터링 체계와 Custom Metric을 활용한 오토스케일링으로 운영 병목을 사전에 대응하는 구조를 구축했습니다.

Apple/Google IAP 결제 라이프사이클을 비동기 웹훅과 상태 머신으로 통합 설계한 경험이 있으며, 트랜잭션 안정성과 동시성을 고려한 설계를 중요하게 생각합니다. 자동화 파이프라인으로 운영 효율 27% 향상, 연간 40일 절감 등 생산성에 지속적으로 기여하며 매 분기 Top Contributor에 선정되었습니다.

CI/CD 파이프라인 구축, Unit Test 고도화, 부하 테스트 설계 가이드 수립 등 배포 안정성과 품질 확보에도 기여해왔으며, 제한된 리소스 속에서도 비즈니스 임팩트를 만들어내는 것을 가장 중요하게 생각합니다."""

p = doc.add_paragraph()
p.space_before = Pt(2)
p.space_after = Pt(2)
run = p.add_run(intro)
run.font.size = Pt(10)


# ============================================================
# 서명
# ============================================================
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(20)
run = p.add_run('상기에 기술한 내용은 사실과 다름 없음을 확인합니다.')
run.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026년    월    일  지원자 정주안 배상.')
run.font.size = Pt(10)


# Save
out_path = os.path.join(os.path.dirname(__file__), '..', 'data', 'references', '이력서_정주안.docx')
doc.save(out_path)
print(f'Saved to {out_path}')
